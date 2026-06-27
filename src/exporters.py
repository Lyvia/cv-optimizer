"""
exporters.py — Convert LLM markdown output to downloadable DOCX files.
Uses python-docx. Handles headings, bullets, bold/italic, paragraphs.
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .styles import StyleConfig, DEFAULT_STYLE


class DOCXExporter:
    """
    Convert markdown-formatted text (as returned by the LLM) to DOCX.

    Design decisions:
    - Clean, minimal styling suitable for HR/corporate contexts
    - ATS-safe: no tables, no text boxes, no headers/footers with content
    - Margins set per document type (CV slightly narrower, letter wider)
    - Colors and font come from a StyleConfig (see styles.py), defaulting
      to the original look if none is provided
    """

    # ── Public ────────────────────────────────────────────────────────────────

    def cv_to_docx(self, markdown_text: str, style: StyleConfig = DEFAULT_STYLE) -> bytes:
        """
        Convert an optimized CV (markdown) to a DOCX byte string.

        Args:
            markdown_text: Markdown-formatted CV content
            style: Visual style (colors + font) to apply

        Returns:
            bytes: DOCX file content ready for st.download_button
        """
        doc = Document()
        self._set_margins(doc, top=0.8, bottom=0.8, left=1.0, right=1.0)
        self._set_default_font(doc, font_name=style.font, font_size=11)
        self._parse_markdown(doc, markdown_text, style)
        return self._to_bytes(doc)

    def cover_letter_to_docx(self, markdown_text: str, style: StyleConfig = DEFAULT_STYLE) -> bytes:
        """
        Convert a cover letter (markdown) to a DOCX byte string.

        Args:
            markdown_text: Markdown-formatted cover letter
            style: Visual style (colors + font) to apply

        Returns:
            bytes: DOCX file content ready for st.download_button
        """
        doc = Document()
        self._set_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.2)
        self._set_default_font(doc, font_name=style.font, font_size=11)
        self._parse_markdown(doc, markdown_text, style)
        return self._to_bytes(doc)

    # ── Core markdown parser ──────────────────────────────────────────────────

    def _parse_markdown(self, doc: Document, text: str, style: StyleConfig):
        """
        Parse a markdown string and populate the Document accordingly.

        Supported elements:
        - # H1 (centered, large)
        - ## H2 (section heading, colored)
        - ### H3 (subsection)
        - - / * bullet points
        - 1. 2. numbered lists
        - **bold** and *italic* inline
        - --- horizontal rule (thin line)
        - blank lines → paragraph spacing
        """
        lines = text.split("\n")
        heading_rgb = RGBColor.from_string(style.heading_color.lstrip("#"))
        accent_rgb = style.accent_color.lstrip("#")

        for line in lines:
            stripped = line.rstrip()

            # ── Empty line
            if not stripped:
                doc.add_paragraph()
                continue

            # ── H1 : # Title
            if stripped.startswith("# ") and not stripped.startswith("## "):
                content = stripped[2:].strip()
                p = doc.add_heading(level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(content)
                run.font.size = Pt(18)
                run.font.color.rgb = heading_rgb
                continue

            # ── H2 : ## Section
            if stripped.startswith("## ") and not stripped.startswith("### "):
                content = stripped[3:].strip()
                if style.heading_uppercase:
                    content = content.upper()
                p = doc.add_heading(level=2)
                run = p.add_run(content)
                run.font.size = Pt(11)
                run.font.color.rgb = heading_rgb
                run.bold = True
                if style.heading_border:
                    self._add_bottom_border(p, accent_rgb)
                continue

            # ── H3 : ### Subsection
            if stripped.startswith("### "):
                content = stripped[4:].strip()
                p = doc.add_heading(level=3)
                run = p.add_run(content)
                run.font.size = Pt(11)
                run.font.color.rgb = heading_rgb
                run.bold = True
                continue

            # ── Horizontal rule
            if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
                p = doc.add_paragraph()
                self._add_bottom_border(p, accent_rgb)
                continue

            # ── Bullet list : - or *
            if re.match(r"^[-*]\s", stripped):
                content = stripped[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                self._add_inline_formatting(p, content, style)
                continue

            # ── Numbered list : 1. 2. etc.
            if re.match(r"^\d+\.\s", stripped):
                content = re.sub(r"^\d+\.\s", "", stripped).strip()
                p = doc.add_paragraph(style="List Number")
                self._add_inline_formatting(p, content, style)
                continue

            # ── Regular paragraph
            p = doc.add_paragraph()
            self._add_inline_formatting(p, stripped, style)

    # ── Inline formatting ─────────────────────────────────────────────────────

    def _add_inline_formatting(self, paragraph, text: str, style: StyleConfig):
        """
        Parse **bold**, *italic*, and ***bold-italic*** inline markers
        and add formatted runs to the paragraph, colored with the main
        body text color from the style.
        """
        text_rgb = RGBColor.from_string(style.text_color.lstrip("#"))

        # Pattern: ***bold-italic***, **bold**, *italic*, plain text
        pattern = re.compile(r"(\*{3}.+?\*{3}|\*{2}.+?\*{2}|\*.+?\*)")
        parts = pattern.split(text)

        for part in parts:
            if part.startswith("***") and part.endswith("***"):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)
            run.font.color.rgb = text_rgb

    # ── Document styling helpers ──────────────────────────────────────────────

    def _set_margins(
        self,
        doc: Document,
        top: float,
        bottom: float,
        left: float,
        right: float,
    ):
        """Set page margins in inches."""
        for section in doc.sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)

    def _set_default_font(self, doc: Document, font_name: str, font_size: int):
        """Set document-level default font."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

    def _add_bottom_border(self, paragraph, color_hex: str):
        """Add a thin bottom border to a paragraph (used for H2 and HR)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _to_bytes(self, doc: Document) -> bytes:
        """Save document to bytes buffer."""
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
