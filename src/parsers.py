"""
parsers.py — Extract plain text from CV and job description files.
Supports: PDF, DOCX, TXT
"""

import io


def parse_document(file) -> str:
    """
    Parse an uploaded Streamlit file object and return its text content.

    Args:
        file: Streamlit UploadedFile object

    Returns:
        str: Extracted text content

    Raises:
        ValueError: If the file format is unsupported or parsing fails
    """
    name = file.name.lower()

    if name.endswith(".pdf"):
        return _parse_pdf(file)
    elif name.endswith(".docx"):
        return _parse_docx(file)
    elif name.endswith(".txt"):
        return _parse_txt(file)
    else:
        raise ValueError(
            f"Unsupported format: '{file.name}'. "
            "Accepted formats: PDF, DOCX, TXT."
        )


# ─── Private parsers ──────────────────────────────────────────────────────────

def _parse_pdf(file) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        )

    try:
        # pdfplumber accepts file-like objects
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(text.strip())

        if not pages_text:
            raise ValueError(
                "No text extracted from the PDF. "
                "The file might be a scanned image (no OCR)."
            )

        return "\n\n".join(pages_text)

    except Exception as e:
        if "pdfplumber" in str(type(e).__module__):
            raise ValueError(f"Error reading PDF: {e}")
        raise


def _parse_docx(file) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    try:
        file.seek(0)
        doc = Document(io.BytesIO(file.read()))

        parts = []

        # Main body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables (for CV tables)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append(" | ".join(row_cells))

        if not parts:
            raise ValueError("No text found in the DOCX file.")

        return "\n".join(parts)

    except Exception as e:
        raise ValueError(f"Error reading DOCX: {e}")


def _parse_txt(file) -> str:
    """Read a plain text file."""
    try:
        file.seek(0)
        content = file.read()
        # Try UTF-8 first, fall back to latin-1
        try:
            return content.decode("utf-8").strip()
        except UnicodeDecodeError:
            return content.decode("latin-1").strip()
    except Exception as e:
        raise ValueError(f"Error reading TXT: {e}")
